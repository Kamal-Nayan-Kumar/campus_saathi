"""Ingestion pipeline: dual parse (Firecrawl + local) -> pick best ->
LangChain RecursiveCharacterTextSplitter -> Knowledge Base upsert.
Small chunks (1000/150) kept. One upload = one ingestion run.
Supports PDF (Firecrawl + pypdf), Excel (Firecrawl + openpyxl),
Word (.docx via python-docx) and plain text (.txt, .md).
"""

import io
import os

from firecrawl import Firecrawl
from firecrawl.v2.types import ParseOptions
from langchain_text_splitters import RecursiveCharacterTextSplitter


class PDFProcessor:
    def __init__(self, knowledge_base):
        api_key = os.getenv("FIRECRAWL_API_KEY")
        if not api_key:
            raise ValueError("FIRECRAWL_API_KEY not found in environment variables")
        self.firecrawl = Firecrawl(api_key=api_key)
        self.knowledge_base = knowledge_base
        self.splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)

    # --- public ---

    def process_and_ingest(self, file_bytes: bytes, filename: str) -> int:
        print(f"Processing {filename}...")
        lower = filename.lower()

        if lower.endswith(".pdf"):
            candidates = self._extract_pdf_both(file_bytes, filename)
        elif lower.endswith((".xlsx", ".xls")):
            candidates = self._extract_excel_both(file_bytes, filename)
        elif lower.endswith(".docx"):
            candidates = self._extract_docx(file_bytes, filename)
        elif lower.endswith((".doc",)):
            raise ValueError("Old .doc format is not supported — please save as .docx")
        elif lower.endswith((".txt", ".md")):
            candidates = self._extract_text(file_bytes, filename)
        else:
            raise ValueError("Unsupported file type")

        # Pick best candidate by simple quality score
        best_text, best_source = self._pick_best(candidates, filename)
        print(f"Picked source: {best_source} ({len(best_text)} chars)")

        if not best_text or not best_text.strip():
            raise ValueError(f"Failed to extract content from {filename}")

        chunks = self.splitter.split_text(best_text)
        print(f"Split document into {len(chunks)} chunks (small chunks kept).")
        count = self.knowledge_base.upsert_chunks(filename, chunks)
        print(f"Ingested {count} chunks from {filename} into the Knowledge Base.")
        return count

    # --- PDF: both parsers ---

    def _extract_pdf_both(self, file_bytes: bytes, filename: str) -> dict[str, str]:
        candidates: dict[str, str] = {}

        # 1) Firecrawl
        try:
            doc = self.firecrawl.parse(
                file_bytes,
                filename=filename,
                content_type="application/pdf",
                options=ParseOptions(formats=["markdown"]),
            )
            candidates["firecrawl"] = (doc.markdown or "").strip()
            print(f"  Firecrawl: {len(candidates['firecrawl'])} chars, has &amp;: {'&amp;' in candidates['firecrawl']}")
        except Exception as e:
            print(f"  Firecrawl failed: {e}")
            candidates["firecrawl"] = ""

        # 2) pypdf (local, no network, keeps table headers like 'For 3rd, 5th and 7th Semester')
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            parts: list[str] = []
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t.strip())
            candidates["pypdf"] = "\n\n".join(parts).strip()
            print(f"  pypdf: {len(candidates['pypdf'])} chars")
        except Exception as e:
            print(f"  pypdf failed: {e}")
            candidates["pypdf"] = ""

        return candidates

    # --- Excel: both parsers ---

    def _extract_excel_both(self, file_bytes: bytes, filename: str) -> dict[str, str]:
        candidates: dict[str, str] = {}

        # 1) openpyxl (local, fast, deterministic for mess menu) — try first
        try:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
            parts: list[str] = []
            for ws in wb.worksheets:
                parts.append(f"# Sheet: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() if c is not None else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
                parts.append("")
            candidates["openpyxl"] = "\n".join(parts).strip()
            print(f"  openpyxl: {len(candidates['openpyxl'])} chars")
        except Exception as e:
            print(f"  openpyxl failed: {e}")
            candidates["openpyxl"] = ""

        # 2) Firecrawl (try with short timeout — fallback only if openpyxl empty)
        if not candidates.get("openpyxl"):
            try:
                doc = self.firecrawl.parse(
                    file_bytes,
                    filename=filename,
                    content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    options=ParseOptions(formats=["markdown"]),
                )
                candidates["firecrawl"] = (doc.markdown or "").strip()
                print(f"  Firecrawl (xlsx): {len(candidates['firecrawl'])} chars")
            except Exception as e:
                print(f"  Firecrawl (xlsx) failed: {e}")
                candidates["firecrawl"] = ""
        else:
            candidates["firecrawl"] = ""

        return candidates

    # --- Text / Markdown: plain decode ---
    def _extract_text(self, file_bytes: bytes, filename: str) -> dict[str, str]:
        candidates: dict[str, str] = {}
        # Try UTF-8 then latin-1
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                text = file_bytes.decode(enc)
                candidates["text"] = text.strip()
                print(f"  text ({enc}): {len(candidates['text'])} chars")
                break
            except Exception:
                continue
        if "text" not in candidates:
            candidates["text"] = file_bytes.decode("utf-8", errors="ignore").strip()
        return candidates

    # --- Word docx ---
    def _extract_docx(self, file_bytes: bytes, filename: str) -> dict[str, str]:
        candidates: dict[str, str] = {}
        # 1) local python-docx
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            # tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            candidates["docx"] = "\n".join(parts).strip()
            print(f"  docx: {len(candidates['docx'])} chars")
        except Exception as e:
            print(f"  docx failed: {e}")
            candidates["docx"] = ""
        # 2) Firecrawl fallback if local empty
        if not candidates.get("docx"):
            try:
                doc = self.firecrawl.parse(
                    file_bytes,
                    filename=filename,
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    options=ParseOptions(formats=["markdown"]),
                )
                candidates["firecrawl"] = (doc.markdown or "").strip()
                print(f"  Firecrawl (docx): {len(candidates['firecrawl'])} chars")
            except Exception as e:
                print(f"  Firecrawl (docx) failed: {e}")
                candidates["firecrawl"] = ""
        else:
            candidates["firecrawl"] = ""
        return candidates

    # --- selector ---

    def _pick_best(self, candidates: dict[str, str], filename: str) -> tuple[str, str]:
        def score(text: str) -> int:
            if not text:
                return -1
            s = len(text)
            s -= text.count("&amp;") * 50  # corrupted tables
            s -= text.count("OnlineOnline") * 20  # Firecrawl concatenation artefact
            if "INDIAN INSTITUTE" in text:
                s += 200  # keeps semester header context (calendar)
            if "For 3rd" in text and "Semester" in text:
                s += 300  # critical header for calendar
            if "Fee Structure" in text or "Tuition Fee" in text:
                s += 100
            if "Mess" in text or "Menu" in text or "Breakfast" in text:
                s += 100
            return s

        best_source = max(candidates, key=lambda k: score(candidates[k]))
        best_text = candidates[best_source]

        # Log scores for transparency
        for src, txt in candidates.items():
            print(f"    score[{src}]={score(txt)} ({len(txt)} chars)")

        if not best_text:
            # fallback to any non-empty
            for src, txt in candidates.items():
                if txt.strip():
                    return txt, src
            return "", "none"
        return best_text, best_source
